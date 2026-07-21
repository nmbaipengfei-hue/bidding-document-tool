"""
排版引擎 - 处理文档排版和美观度
"""
import logging
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class LayoutEngine:
    """文档排版引擎"""
    
    def __init__(self):
        """初始化排版引擎"""
        self.default_font_name = '微软雅黑'
        self.default_font_size = 11
        self.colors = {
            'primary': RGBColor(0, 51, 102),
            'secondary': RGBColor(102, 102, 102),
            'accent': RGBColor(255, 102, 0),
        }
    
    def set_page_margin(self, document, top=1.27, bottom=1.27, left=1.27, right=1.27):
        """
        设置页边距
        
        Args:
            document: Word 文档对象
            top: 上边距 (厘米)
            bottom: 下边距 (厘米)
            left: 左边距 (厘米)
            right: 右边距 (厘米)
        """
        try:
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(top / 2.54)
                section.bottom_margin = Inches(bottom / 2.54)
                section.left_margin = Inches(left / 2.54)
                section.right_margin = Inches(right / 2.54)
            logger.info("页边距已设置")
        except Exception as e:
            logger.error(f"设置页边距失败: {str(e)}")
    
    def set_font(self, paragraph, font_name=None, font_size=None, bold=False, 
                 italic=False, color=None):
        """
        设置段落字体
        
        Args:
            paragraph: 段落对象
            font_name: 字体名称
            font_size: 字体大小
            bold: 是否加粗
            italic: 是否斜体
            color: 字体颜色
        """
        try:
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
            font = run.font
            
            if font_name:
                font.name = font_name
            if font_size:
                font.size = Pt(font_size)
            if bold:
                font.bold = True
            if italic:
                font.italic = True
            if color:
                font.color.rgb = color
        except Exception as e:
            logger.error(f"设置字体失败: {str(e)}")
    
    def set_paragraph_style(self, paragraph, alignment=None, space_before=None, 
                           space_after=None, line_spacing=None):
        """
        设置段落样式
        
        Args:
            paragraph: 段落对象
            alignment: 对齐方式
            space_before: 段前距离
            space_after: 段后距离
            line_spacing: 行距
        """
        try:
            if alignment:
                paragraph.alignment = alignment
            
            pPr = paragraph._element.get_or_add_pPr()
            
            if space_before is not None:
                pSpacing = pPr.find(qn('w:spacing'))
                if pSpacing is None:
                    pSpacing = OxmlElement('w:spacing')
                    pPr.append(pSpacing)
                pSpacing.set(qn('w:before'), str(int(space_before * 20)))
            
            if space_after is not None:
                pSpacing = pPr.find(qn('w:spacing'))
                if pSpacing is None:
                    pSpacing = OxmlElement('w:spacing')
                    pPr.append(pSpacing)
                pSpacing.set(qn('w:after'), str(int(space_after * 20)))
            
            if line_spacing is not None:
                pSpacing = pPr.find(qn('w:spacing'))
                if pSpacing is None:
                    pSpacing = OxmlElement('w:spacing')
                    pPr.append(pSpacing)
                pSpacing.set(qn('w:line'), str(int(line_spacing * 240)))
        except Exception as e:
            logger.error(f"设置段落样式失败: {str(e)}")
    
    def add_heading_style(self, paragraph, level=1):
        """
        添加标题样式
        
        Args:
            paragraph: 段落对象
            level: 标题级别
        """
        try:
            paragraph.style = f'Heading {level}'
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if paragraph.runs:
                if level == 1:
                    self.set_font(paragraph, font_size=16, bold=True, 
                                color=self.colors['primary'])
                elif level == 2:
                    self.set_font(paragraph, font_size=14, bold=True, 
                                color=self.colors['secondary'])
                else:
                    self.set_font(paragraph, font_size=12, bold=True)
        except Exception as e:
            logger.error(f"添加标题样式失败: {str(e)}")
    
    def style_table(self, table, style_name='Light Grid Accent 1'):
        """
        设置表格样式
        
        Args:
            table: 表格对象
            style_name: 表格样式名称
        """
        try:
            table.style = style_name
            
            # 设置表头样式
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        except Exception as e:
            logger.error(f"设置表格样式失败: {str(e)}")
    
    def set_column_width(self, table, column_index, width):
        """
        设置表格列宽
        
        Args:
            table: 表格对象
            column_index: 列索引
            width: 列宽 (厘米)
        """
        try:
            for row in table.rows:
                row.cells[column_index].width = Inches(width / 2.54)
        except Exception as e:
            logger.error(f"设置列宽失败: {str(e)}")
    
    def set_row_height(self, table, row_index, height):
        """
        设置表格行高
        
        Args:
            table: 表格对象
            row_index: 行索引
            height: 行高 (厘米)
        """
        try:
            table.rows[row_index].height = Inches(height / 2.54)
        except Exception as e:
            logger.error(f"设置行高失败: {str(e)}")
    
    def add_border(self, table):
        """
        为表格添加边框
        
        Args:
            table: 表格对象
        """
        try:
            tbl = table._element
            tblPr = tbl.tblPr
            
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # 创建表格边框
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
        except Exception as e:
            logger.error(f"添加边框失败: {str(e)}")
    
    def optimize_layout(self, document):
        """
        优化文档排版
        
        Args:
            document: Word 文档对象
        """
        try:
            # 设置默认页边距
            self.set_page_margin(document)
            
            # 设置段落间距和行距
            for paragraph in document.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    self.set_paragraph_style(paragraph, 
                                           space_before=0.5, 
                                           space_after=0.5,
                                           line_spacing=1.5)
                else:
                    self.set_paragraph_style(paragraph, 
                                           space_after=0.2,
                                           line_spacing=1.15)
            
            logger.info("文档排版已优化")
        except Exception as e:
            logger.error(f"优化排版失败: {str(e)}")