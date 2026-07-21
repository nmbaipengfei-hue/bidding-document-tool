"""
模板管理模块
"""
import logging
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class TemplateManager:
    """标书模板管理器"""
    
    def __init__(self, template_dir=None):
        """
        初始化模板管理器
        
        Args:
            template_dir: 模板目录路径
        """
        self.template_dir = Path(template_dir) if template_dir else Path('templates')
        self.templates = {}
        self.current_document = None
        
    def create_blank_template(self, title="标书"):
        """
        创建空白模板
        
        Args:
            title: 文档标题
            
        Returns:
            Document: Word 文档对象
        """
        try:
            doc = Document()
            
            # 设置文档样式
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style.font.size = Pt(11)
            
            # 添加标题
            title_para = doc.add_paragraph(title, style='Heading 1')
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            self.current_document = doc
            logger.info(f"创建空白模板: {title}")
            return doc
        except Exception as e:
            logger.error(f"创建空白模板失败: {str(e)}")
            return None
    
    def load_template(self, template_name):
        """
        加载模板
        
        Args:
            template_name: 模板名称
            
        Returns:
            Document: Word 文档对象
        """
        try:
            template_path = self.template_dir / f"{template_name}.docx"
            if not template_path.exists():
                logger.warning(f"模板不存在: {template_path}")
                return self.create_blank_template(template_name)
            
            doc = Document(template_path)
            self.current_document = doc
            logger.info(f"加载模板: {template_name}")
            return doc
        except Exception as e:
            logger.error(f"加载模板失败: {str(e)}")
            return None
    
    def save_template(self, template_name, document=None):
        """
        保存模板
        
        Args:
            template_name: 模板名称
            document: Word 文档对象
            
        Returns:
            bool: 保存是否成功
        """
        try:
            doc = document or self.current_document
            if not doc:
                return False
            
            self.template_dir.mkdir(parents=True, exist_ok=True)
            template_path = self.template_dir / f"{template_name}.docx"
            doc.save(template_path)
            
            logger.info(f"模板已保存: {template_path}")
            return True
        except Exception as e:
            logger.error(f"保存模板失败: {str(e)}")
            return False
    
    def add_heading(self, text, level=1, document=None):
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别
            document: Word 文档对象
        """
        doc = document or self.current_document
        if doc:
            para = doc.add_paragraph(text, style=f'Heading {level}')
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def add_paragraph(self, text, style='Normal', document=None):
        """
        添加段落
        
        Args:
            text: 段落文本
            style: 段落样式
            document: Word 文档对象
        """
        doc = document or self.current_document
        if doc:
            doc.add_paragraph(text, style=style)
    
    def add_table(self, rows, cols, document=None):
        """
        添加表格
        
        Args:
            rows: 行数
            cols: 列数
            document: Word 文档对象
            
        Returns:
            Table: 表格对象
        """
        doc = document or self.current_document
        if doc:
            return doc.add_table(rows=rows, cols=cols)
        return None
    
    def add_image(self, image_path, width=None, height=None, document=None):
        """
        添加图像
        
        Args:
            image_path: 图像路径
            width: 图像宽度 (英寸)
            height: 图像高度 (英寸)
            document: Word 文档对象
        """
        try:
            doc = document or self.current_document
            if doc:
                if width and height:
                    doc.add_picture(str(image_path), width=Inches(width), height=Inches(height))
                elif width:
                    doc.add_picture(str(image_path), width=Inches(width))
                else:
                    doc.add_picture(str(image_path))
        except Exception as e:
            logger.error(f"添加图像失败: {str(e)}")
    
    def add_page_break(self, document=None):
        """
        添加分页符
        
        Args:
            document: Word 文档对象
        """
        doc = document or self.current_document
        if doc:
            doc.add_page_break()
    
    def save_document(self, output_path, document=None):
        """
        保存文档
        
        Args:
            output_path: 输出路径
            document: Word 文档对象
            
        Returns:
            bool: 保存是否成功
        """
        try:
            doc = document or self.current_document
            if not doc:
                return False
            
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            logger.info(f"文档已保存: {output_path}")
            return True
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            return False
    
    def list_templates(self):
        """
        列出所有可用的模板
        
        Returns:
            list: 模板名称列表
        """
        if not self.template_dir.exists():
            return []
        
        templates = [f.stem for f in self.template_dir.glob('*.docx')]
        return templates