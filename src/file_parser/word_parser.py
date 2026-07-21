"""
Word 文件解析模块
"""
import logging
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


class WordParser:
    """Word 文件解析器"""
    
    def __init__(self, file_path):
        """
        初始化 Word 解析器
        
        Args:
            file_path: Word 文件路径
        """
        self.file_path = Path(file_path)
        self.document = None
        self.content = []
        
    def parse(self):
        """
        解析 Word 文件
        
        Returns:
            dict: 解析结果，包含段落、表格等内容
        """
        try:
            self.document = Document(self.file_path)
            
            self.content = {
                'paragraphs': [],
                'tables': [],
                'images': []
            }
            
            # 解析段落
            for para in self.document.paragraphs:
                if para.text.strip():
                    self.content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name,
                        'level': para.paragraph_format.outline_level
                    })
            
            # 解析表格
            for table_idx, table in enumerate(self.document.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                self.content['tables'].append({
                    'index': table_idx,
                    'data': table_data
                })
            
            logger.info(f"成功解析 Word 文件: {self.file_path}")
            return {
                'success': True,
                'file': str(self.file_path),
                'paragraphs': len(self.content['paragraphs']),
                'tables': len(self.content['tables']),
                'content': self.content
            }
        except Exception as e:
            logger.error(f"Word 解析失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def get_text(self):
        """
        获取提取的文本内容
        
        Returns:
            str: 合并后的文本内容
        """
        return '\n'.join([p['text'] for p in self.content['paragraphs']])
    
    def get_tables(self):
        """
        获取表格数据
        
        Returns:
            list: 表格数据列表
        """
        return self.content['tables']
    
    def extract_keywords(self, keywords):
        """
        从 Word 中提取关键词相关内容
        
        Args:
            keywords: 关键词列表
            
        Returns:
            dict: 关键词及其所在的段落
        """
        results = {}
        
        for keyword in keywords:
            for para in self.content['paragraphs']:
                if keyword.lower() in para['text'].lower():
                    if keyword not in results:
                        results[keyword] = []
                    results[keyword].append(para['text'])
        
        return results